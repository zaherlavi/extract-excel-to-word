import openpyxl
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def write_columns_to_word(file_path, sheet_name, end_column, output_doc):

    workbook = openpyxl.load_workbook(file_path)
    sheet = workbook[sheet_name]
    doc = Document()

    for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, min_col=1, max_col=end_column):
        doc.add_paragraph()
        p = doc.add_paragraph()

        for index, cell in enumerate(row):
            cell_value = str(cell.value)
            
            if index > 0 and index != 1:
                p.add_run('\n')
            if index == 1:
                continue
            if index == 2:
                p.add_run('Global ID: ')
            if index == 5:
                p.add_run('DCS Code: ')

            run = p.add_run(cell_value)

            # Check for highlighting
            if index == 0:
                if row[1].value == 'Green':
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                elif row[1].value == 'Red':
                    run.font.highlight_color = WD_COLOR_INDEX.RED
                elif row[1].value == 'Yellow':
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        doc.add_paragraph()

    doc.save(output_doc)


file_path = 'new.xlsx'
sheet_name = 'Export'
end_column = 6
output_document = f'{file_path[:-5]}.docx' # Output file with its original name without its extenstion

write_columns_to_word(file_path, sheet_name, end_column, output_document)
